import random
from docx import Document

def split_variants(words):
    random.shuffle(words)
    mid = len(words)//2
    return words[:mid], words[mid:]

def generate_doc(words, level, variant):
    doc = Document()

    doc.add_heading(f"Quiz {variant} ({level})", 0)

    # A1–A2 → translation
    if level in ["A1", "A2"]:
        doc.add_heading("Translate into English", 1)
        for i, (_, rus) in enumerate(words):
            doc.add_paragraph(f"{i+1}. {rus} — ______")

    # Writing
    doc.add_heading("Writing", 1)
    word_bank = ", ".join([w[0] for w in words])
    doc.add_paragraph(f"Use 3–5 words: {word_bank}")
    doc.add_paragraph("How can people improve their lifestyle?")

    # Speaking
    doc.add_heading("Speaking", 1)
    doc.add_paragraph("Use at least 2 words from the list.")
    doc.add_paragraph("Why is this topic important today?")

    # Answer key
    doc.add_page_break()
    doc.add_heading("Answer Key", 0)

    for i, (eng, rus) in enumerate(words):
        doc.add_paragraph(f"{i+1}. {rus} → {eng}")

    return doc
